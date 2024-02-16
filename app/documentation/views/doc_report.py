import io
import locale
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

from app.documentation.views.utils import formatear_fecha_esp


# Función para suplantar datos faltantes
def suplantar_dato(dato):
    return dato if dato else ". . . . . . . . ."

def _set_styles(document):
    normal_style = document.styles.add_style('Normal_gris', WD_STYLE_TYPE.PARAGRAPH)
    normal_style.font.name = 'Assistant'
    normal_style.font.size = Pt(12)
    normal_style.font.color.rgb = RGBColor(67, 67, 67)  # Define el color en gris

    titulo_gris_style = document.styles.add_style('Titulo_gris', WD_STYLE_TYPE.PARAGRAPH)
    titulo_gris_style.font.name = 'Assistant'
    titulo_gris_style.font.size = Pt(12)
    titulo_gris_style.font.color.rgb = RGBColor(67, 67, 67)  # Define el color en gris
    titulo_gris_style.font.bold = True


    titulo_style = document.styles.add_style('Title_Custom', WD_STYLE_TYPE.PARAGRAPH)
    titulo_style.font.name = 'Assistant'
    titulo_style.font.size = Pt(16)
    titulo_style.font.bold = True

    bold_style = document.styles.add_style('Bold_Custom', WD_STYLE_TYPE.PARAGRAPH)
    bold_style.font.name = 'Assistant'
    bold_style.font.size = Pt(12)
    bold_style.font.bold = True
    
    return document

def _add_title_report(document, data):
        # Crear una fecha
    fecha_actual = datetime.now()
    fecha_formateada = formatear_fecha_esp(fecha_actual)
    
    # Agregar Título
    txt_titulo = f'Informe Legal - Leasing Retail | {fecha_formateada}, Santiago de Chile'
    document.add_paragraph(txt_titulo, style='Titulo_gris').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Agregar Nombre/Razón Social
    txt_razon_social = f"{suplantar_dato(data.get('nombre_razon_social'))} ({suplantar_dato(data.get('nombre_fantasia'))})"
    document.add_paragraph(txt_razon_social, style='Title_Custom').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Agregar RUT
    txt_rut = f"RUT: {suplantar_dato(data.get('rut'))}"
    document.add_paragraph(txt_rut, style='Bold_Custom').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_paragraph()
    return document

def _add_resultado_report(document, data):
    # Agregar Conclusiones
    p_apto_leasing = document.add_paragraph()
    run = p_apto_leasing.add_run("Resultado: ")
    font = run.font
    font.name = 'Assistant'
    font.size = Pt(12)

    # Agregar Apto/No Apto
    comentario_conclusion = data.get('comentario_conclusion', {})
    apto_aux = comentario_conclusion.get('apto_cliente_leasing')    
    print("DEBUG, APTO_AUX:\n", apto_aux)

    if not apto_aux:
        color_aux = RGBColor(60, 120, 216)  # Define el color en azul
    elif apto_aux.upper() == "APTO":
        color_aux = RGBColor(106, 168, 79)  # Define el color en verde
        print("DEBUG, APTO:\n", color_aux)
    elif apto_aux.upper() == "NO APTO":
        color_aux = RGBColor(204, 65, 37) # Define el color en rojo
        print("DEBUG, NO APTO:\n", color_aux)

    run = p_apto_leasing.add_run(apto_aux)
    font = run.font
    font.name = 'Assistant'
    font.size = Pt(12)
    font.bold = True
    font.color.rgb = color_aux

    # Agregar Comentario conclusion
    document.add_paragraph(comentario_conclusion.get('comentario', ""), style='Normal_gris')
    document.add_paragraph()
    return document

def _add_constitucion_legal_report(document, data):
    document.add_paragraph('Constitución Legal:', style='Bold_Custom')

    # Agregar Registro
    document.add_paragraph(f"Registro: {suplantar_dato(data['constitucion_legal'].get('registro'))}", style='Normal_gris')

    # Agregar Codigo
    document.add_paragraph(f"Código de consulta: {suplantar_dato(data['constitucion_legal'].get('codigo'))}", style='Normal_gris')

    # Agregar Tipo de Sociedad
    document.add_paragraph(f"Tipo de sociedad: {suplantar_dato(data.get('tipo_sociedad'))}", style='Normal_gris')

    # Agregar Domicilio
    document.add_paragraph(f"Domicilio: {suplantar_dato(data.get('domicilio'))}", style='Normal_gris')

    # Agregar Duración
    document.add_paragraph(f"Duración: {suplantar_dato(data.get('plazo'))}", style='Normal_gris')

    # Agregar Objeto Principal
    document.add_paragraph(f"Objeto Principal: {suplantar_dato(data.get('objeto_principal'))}", style='Normal_gris')

    document.add_paragraph()
    return document

def _add_personeria_report(document, data):
    document.add_paragraph('Personería:', style='Bold_Custom')
    document.add_paragraph(data.get('personeria', ""), style='Normal_gris')
    document.add_paragraph()
    return document

def _add_modificaciones_report(document, data):
    
    document.add_paragraph('Modificaciones:', style='Bold_Custom')

    # Verificar si hay modificaciones
    list_modificaciones = data.get('modificacion', [])
    if list_modificaciones:
        modificacion_str = ""
        for modificacion in list_modificaciones:
            mod_desc = modificacion.get('descripcion')
            mod_cod = modificacion.get('codigo', "---")
            mod_fecha = modificacion.get('fecha')            
            modificacion_str += f"{mod_desc} ({mod_cod}) - {mod_fecha}\n"
            
        document.add_paragraph(modificacion_str, style='Normal_gris')
    else:
        document.add_paragraph('No se han encontrado modificaciones.', style='Normal_gris')
    return document

def _add_capital_report(document, data):
    document.add_paragraph('Capital:', style='Bold_Custom')
    document.add_paragraph(suplantar_dato(data.get('capital')), style='Normal_gris')

    document.add_paragraph()
    return document

def _add_accionistas_report(document, data):
    document.add_paragraph('Accionistas:', style='Bold_Custom')
    table = document.add_table(rows=1, cols=3)
    table.autofit = True
    table.style.font.name = 'Assistant'
    table.style.font.size = Pt(12)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nombre'
    hdr_cells[1].text = 'RUT'
    hdr_cells[2].text = 'Aportes'
    
    for socio in data['socios']:
        row_cells = table.add_row().cells
        row_cells[0].text = suplantar_dato(socio.get('nombre_socio'))
        row_cells[1].text = suplantar_dato(socio.get('rut_socio'))
        row_cells[2].text = suplantar_dato(socio.get('aportes_socio'))

    document.add_paragraph()
    return document

def _add_facultades_otorgadas_report(document, data):
    document.add_paragraph('Facultades Otorgadas:', style='Bold_Custom')
    table = document.add_table(rows=1, cols=2)
    table.autofit = True
    table.style.font.name = 'Assistant'
    table.style.font.size = Pt(12)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Facultad'
    hdr_cells[1].text = 'Estado'
    
    for facultad, valor in data.get('facultades_otorgadas', {}).items():
        row_cells = table.add_row().cells
        row_cells[0].text = facultad
        row_cells[1].text = "SI" if valor else "NO"
    
    document.add_paragraph()    
    return document

def build_report(data):
    """
    Genera un informe sobre una sociedad y crea un archivo DOCX con la información proporcionada.
    """
    print("DEBUG, DATA:\n", data)
    # Crea un nuevo documento
    document = Document()

    document = _set_styles(document)
    document = _add_title_report(document, data)
    document = _add_resultado_report(document, data)
    document = _add_constitucion_legal_report(document, data)
    document = _add_personeria_report(document, data)
    document = _add_modificaciones_report(document, data)
    document = _add_capital_report(document, data)
    document = _add_accionistas_report(document, data)
    document = _add_facultades_otorgadas_report(document, data)

    # Guarda el archivo DOCX
    docx_bytes = io.BytesIO()
    document.save(docx_bytes)
    docx_bytes.seek(0)

    return docx_bytes
